# GeneraDenuncia.py
# Generatore di documenti di denuncia per detenzione armi

import sqlite3
from datetime import datetime
import os
import sys
import traceback
import subprocess
import logging
from typing import List, Dict, Any, Optional, Tuple
from collections import defaultdict

# Importa i widget PyQt5 necessari per i dialoghi
from PyQt5.QtWidgets import QMessageBox, QFileDialog, QDialog, QProgressDialog
from PyQt5.QtCore import Qt

# Importa la libreria per gestire i template Word
from docxtpl import DocxTemplate
import jinja2

# Configurazione del logging
logging.basicConfig(
    level=logging.DEBUG,  # Livello DEBUG per log più dettagliati
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler("denunce.log"),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger("GeneraDenuncia")


def get_detentore_data(cursor, detentore_id: int) -> Optional[Dict[str, Any]]:
    """
    Recupera i dati completi del detentore dal database.

    Args:
        cursor: Cursore del database
        detentore_id: ID del detentore

    Returns:
        Dizionario con i dati del detentore o None se non trovato
    """
    try:
        query_detentore = """
            SELECT 
                Cognome, Nome, CodiceFiscale, DataNascita, LuogoNascita, SiglaProvinciaNascita,
                ComuneResidenza, SiglaProvinciaResidenza, TipoVia, Via, Civico, Telefono,
                TipologiaTitolo, NumeroPortoArmi, DataRilascio, EnteRilascio,
                ComuneDetenzione, TipoViaDetenzione, ViaDetenzione, CivicoDetenzione,
                Sesso, SiglaProvinciaDetenzione
            FROM detentori 
            WHERE ID_Detentore = ?
        """
        logger.debug(f"Esecuzione query detentore con ID: {detentore_id}")
        cursor.execute(query_detentore, (detentore_id,))
        row = cursor.fetchone()

        if not row:
            logger.warning(f"Nessun detentore trovato con ID: {detentore_id}")
            return None

        logger.debug(f"Detentore trovato: {row[1]} {row[0]}")
        return {
            'cognome': row[0] or "",
            'nome': row[1] or "",
            'codice_fiscale': row[2] or "",
            'data_nascita': row[3] or "",
            'luogo_nascita': row[4] or "",
            'provincia_nascita': row[5] or "",
            'comune_residenza': row[6] or "",
            'provincia_residenza': row[7] or "",
            'tipo_via': row[8] or "",
            'via': row[9] or "",
            'civico': row[10] or "",
            'telefono': row[11] or "",
            'tipologia_titolo': row[12] or "NESSUN TITOLO",
            'numero_porto_armi': row[13] or "N/D",
            'data_rilascio': row[14] or "N/D",
            'ente_rilascio': row[15] or "N/D",
            'comune_detenzione': row[16] or "",
            'tipo_via_detenzione': row[17] or "",
            'via_detenzione': row[18] or "",
            'civico_detenzione': row[19] or "",
            'sesso': row[20] or "M",
            'provincia_detenzione': row[21] or ""
        }
    except Exception as e:
        logger.error(f"Errore durante il recupero dei dati del detentore: {e}")
        logger.error(traceback.format_exc())
        return None


def get_armi_data(cursor, detentore_id: int) -> Tuple[List[Dict[str, Any]], Dict[str, List[Dict[str, Any]]]]:
    """
    Recupera le armi associate al detentore e le raggruppa per categoria.

    Args:
        cursor: Cursore del database
        detentore_id: ID del detentore

    Returns:
        Tuple(lista completa di armi, dizionario con armi raggruppate per categoria)
    """
    try:
        # Prima verifichiamo quante armi ha il detentore con una query semplice
        cursor.execute("SELECT COUNT(*) FROM armi WHERE ID_Detentore = ?", (detentore_id,))
        count = cursor.fetchone()[0]
        logger.debug(f"Il detentore con ID {detentore_id} ha {count} armi nel database")

        # Query completa per recuperare tutti i dati delle armi
        query_armi = """
            SELECT 
                ID_ArmaDetenuta, TipoArma, MarcaArma, ModelloArma, Matricola, 
                CalibroArma, NoteArma, CategoriaArma, LunghezzaCanna, NumeroCanne,
                FunzionamentoArma, CaricamentoArma, ArmaLungaCorta, TipologiaArma
            FROM armi 
            WHERE ID_Detentore = ?
            ORDER BY CategoriaArma, MarcaArma, ModelloArma
        """
        logger.debug(f"Esecuzione query armi per detentore ID: {detentore_id}")
        logger.debug(f"Query: {query_armi}")

        cursor.execute(query_armi, (detentore_id,))
        armi_db = cursor.fetchall()

        logger.debug(f"Numero di armi trovate: {len(armi_db)}")

        if not armi_db:
            logger.warning(f"Nessuna arma trovata per il detentore con ID: {detentore_id}")
            return [], {}

        armi_formattate = []
        armi_per_categoria = defaultdict(list)

        categorie_standard = {
            "": "ALTRA CATEGORIA",
            "N/D": "ALTRA CATEGORIA",
            None: "ALTRA CATEGORIA"
        }

        # Mappa delle categorie per ordinamento
        categoria_ordine = {
            "ARMA DA CACCIA": 1,
            "ARMA SPORTIVA": 2,
            "ARMA COMUNE": 3,
            "ARMA ANTICA": 4,
            "ALTRA CATEGORIA": 99
        }

        for idx, arma in enumerate(armi_db):
            logger.debug(f"Elaborazione arma {idx + 1}: {arma}")

            id_arma = arma[0] if len(arma) > 0 else "N/D"
            tipo_arma = arma[1] or "N/D"
            marca = arma[2] or "N/D"
            modello = arma[3] or "N/D"
            matricola = arma[4] or "N/D"
            calibro = arma[5] or "N/D"
            note = arma[6] or ""
            categoria = arma[7] or "ALTRA CATEGORIA"
            lung_canna = arma[8] or "N/D"
            num_canne = arma[9] or "1"
            funzionamento = arma[10] or "N/D"
            caricamento = arma[11] or "N/D"
            lunga_corta = arma[12] or "ARMA CORTA"
            tipologia = arma[13] or "N/D"

            # Standardizza la categoria
            categoria = categoria.strip().upper()
            if categoria in categorie_standard:
                categoria = categorie_standard[categoria]

            # Formattazione speciale per il documento
            descrizione_completa = f"{tipo_arma} {marca} {modello}"
            if calibro and calibro != "N/D":
                descrizione_completa += f" cal. {calibro}"
            if note:
                descrizione_completa += f" - {note}"

            arma_formattata = {
                'id': id_arma,
                'tipo': tipo_arma,
                'marca': marca,
                'modello': modello,
                'matricola': matricola,
                'calibro': calibro,
                'note': note,
                'categoria': categoria,
                'lung_canna': lung_canna,
                'num_canne': num_canne,
                'funzionamento': funzionamento,
                'caricamento': caricamento,
                'lunga_corta': lunga_corta,
                'tipologia': tipologia,
                'descrizione_completa': descrizione_completa,
                'catalogo': "N/D",  # Aggiunto per compatibilità con il template
                'categoria_ordine': categoria_ordine.get(categoria, 99)  # Per ordinamento
            }

            logger.debug(f"Arma formattata: {arma_formattata}")
            armi_formattate.append(arma_formattata)

            # Aggiungi l'arma alla sua categoria nel dizionario
            armi_per_categoria[categoria].append(arma_formattata)

        # Registra le categorie trovate
        logger.debug(f"Categorie di armi trovate: {list(armi_per_categoria.keys())}")

        return armi_formattate, armi_per_categoria
    except Exception as e:
        logger.error(f"Errore durante il recupero delle armi: {e}")
        logger.error(traceback.format_exc())
        return [], {}


def find_template_path(template_name: str) -> Optional[str]:
    """
    Cerca il percorso del template in diverse posizioni.

    Args:
        template_name: Nome del file template

    Returns:
        Percorso completo del template o None se non trovato
    """
    # Possibili percorsi del template
    possible_paths = [
        os.path.join(os.getcwd(), "templates", template_name),
        os.path.join(os.getcwd(), template_name),
        os.path.join(os.path.dirname(sys.executable), "templates", template_name),
        os.path.join(os.path.dirname(sys.executable), template_name),
        os.path.join(os.path.dirname(__file__), "templates", template_name),
        os.path.join(os.path.dirname(__file__), template_name)
    ]

    logger.debug(f"Ricerca template: {template_name}")

    # Verifica se esiste il percorso
    for path in possible_paths:
        logger.debug(f"Verifica percorso: {path}")
        if os.path.exists(path):
            logger.debug(f"Template trovato in: {path}")
            return path

    logger.warning(f"Template '{template_name}' non trovato!")
    return None


def prepare_context(detentore: Dict[str, Any], armi: List[Dict[str, Any]],
                    armi_per_categoria: Dict[str, List[Dict[str, Any]]]) -> Dict[str, Any]:
    """
    Prepara il contesto per il template.

    Args:
        detentore: Dati del detentore
        armi: Lista delle armi
        armi_per_categoria: Dizionario con armi raggruppate per categoria

    Returns:
        Dizionario con il contesto completo per il template
    """
    data_corrente = datetime.now().strftime('%d/%m/%Y')

    # Determina l'indirizzo di detenzione
    indirizzo_detenzione = ""
    if detentore.get('tipo_via_detenzione') and detentore.get('via_detenzione'):
        indirizzo_detenzione = f"{detentore['tipo_via_detenzione']} {detentore['via_detenzione']}"
        if detentore.get('civico_detenzione'):
            indirizzo_detenzione += f" {detentore['civico_detenzione']}"

    # Determina l'indirizzo di residenza
    indirizzo_residenza = ""
    if detentore.get('tipo_via') and detentore.get('via'):
        indirizzo_residenza = f"{detentore['tipo_via']} {detentore['via']}"
        if detentore.get('civico'):
            indirizzo_residenza += f" {detentore['civico']}"

    # Prepara la lista delle categorie con relativi conteggi
    categorie = []
    for categoria, armi_cat in armi_per_categoria.items():
        categorie.append({
            'nome': categoria,
            'count': len(armi_cat),
            'armi': armi_cat
        })

    # Ordina le categorie secondo l'ordine specificato
    categorie_ordinate = sorted(categorie, key=lambda x: (
        1 if x['nome'] == "ARMA DA CACCIA" else
        2 if x['nome'] == "ARMA SPORTIVA" else
        3 if x['nome'] == "ARMA COMUNE" else
        4 if x['nome'] == "ARMA ANTICA" else
        99
    ))

    # Preparazione del contesto completo
    context = {
        'comando_stazione': "COMANDO STAZIONE CARABINIERI",
        'cognome': detentore['cognome'],
        'nome': detentore['nome'],
        'codice_fiscale': detentore['codice_fiscale'],
        'data_nascita': detentore['data_nascita'],
        'luogo_nascita': detentore['luogo_nascita'],
        'provincia_nascita': detentore['provincia_nascita'],
        'comune_residenza': detentore['comune_residenza'],
        'provincia_residenza': detentore['provincia_residenza'],
        'indirizzo_residenza': indirizzo_residenza,
        'tipo_via': detentore['tipo_via'],
        'via': detentore['via'],
        'civico': detentore['civico'],
        'telefono': detentore['telefono'],
        'tipologia_titolo': detentore['tipologia_titolo'],
        'numero_porto_armi': detentore['numero_porto_armi'],
        'data_rilascio': detentore['data_rilascio'],
        'ente_rilascio': detentore['ente_rilascio'],
        'comune_detenzione': detentore['comune_detenzione'],
        'provincia_detenzione': detentore['provincia_detenzione'],
        'tipo_via_detenzione': detentore['tipo_via_detenzione'],
        'via_detenzione': detentore['via_detenzione'],
        'civico_detenzione': detentore['civico_detenzione'],
        'indirizzo_detenzione': indirizzo_detenzione,
        'armi': armi,  # Lista completa per compatibilità
        'categorie': categorie_ordinate,  # Armi raggruppate per categoria
        'data_documento': data_corrente,
        'num_armi': len(armi),
        'num_categorie': len(categorie_ordinate),
        'sesso': detentore['sesso'],
        'user': os.getlogin() if hasattr(os, 'getlogin') else "Utente",  # Per il nuovo template
        'timestamp': datetime.now().strftime('%Y-%m-%d %H:%M:%S')  # Per il nuovo template
    }

    logger.debug(f"Contesto preparato con {len(armi)} armi in {len(categorie_ordinate)} categorie")
    return context


def crea_template_professionale():
    """
    Crea il template professionale se non esiste già.

    Returns:
        str: Percorso del template creato o esistente
    """
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    template_path = "Denuncia_armi_completo.docx"

    # Se il template esiste già, restituisci il percorso
    if os.path.exists(template_path):
        return os.path.abspath(template_path)

    # Altrimenti, crea un nuovo template
    try:
        # Crea un nuovo documento Word
        doc = Document()

        # Imposta margini
        sections = doc.sections
        for section in sections:
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

        # Titolo
        title = doc.add_heading('DENUNCIA DI DETENZIONE ARMI', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Dati personali
        p = doc.add_paragraph()
        p.add_run('Il/La sottoscritto/a ').bold = False
        p.add_run('{{ cognome }} {{ nome }}').bold = True
        p.add_run(', nato/a a ').bold = False
        p.add_run('{{ luogo_nascita }} ({{ provincia_nascita }})').bold = True
        p.add_run(' il ').bold = False
        p.add_run('{{ data_nascita }}').bold = True
        p.add_run(', residente a ').bold = False
        p.add_run('{{ comune_residenza }} ({{ provincia_residenza }})').bold = True

        # Dichiarazione
        doc.add_heading('DICHIARA', level=1)
        dichiarazione = doc.add_paragraph()
        dichiarazione.add_run(
            'ai sensi dell\'articolo 38 del T.U.L.P.S. di detenere presso la propria abitazione le seguenti armi, suddivise per categoria:')

        # Loop per le categorie
        doc.add_paragraph('{% for categoria in categorie %}')

        # Titolo categoria
        cat_title = doc.add_paragraph()
        cat_title.add_run('{{ loop.index }}. CATEGORIA: {{ categoria.nome }} ({{ categoria.count }} armi)').bold = True

        # Loop per le armi di questa categoria
        doc.add_paragraph('{% for arma in categoria.armi %}')

        # Descrizione arma
        p_arma = doc.add_paragraph()
        p_arma.style = 'List Bullet'
        p_arma.add_run(
            '{{ arma.tipo }} {{ arma.marca }} {{ arma.modello }}, matricola {{ arma.matricola }}, calibro {{ arma.calibro }}')

        # Chiudi il loop delle armi
        doc.add_paragraph('{% endfor %}')

        # Chiudi il loop delle categorie
        doc.add_paragraph('{% endfor %}')

        # Data e firma
        doc.add_paragraph()
        data = doc.add_paragraph()
        data.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        data.add_run('{{ comune_residenza }}, {{ data_documento }}')

        doc.add_paragraph()
        doc.add_paragraph()

        firma = doc.add_paragraph()
        firma.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        firma.add_run('Firma').bold = True
        firma.add_run('\n___________________________')

        # Salva il template
        doc.save(template_path)
        logger.info(f"Template base creato con successo: {os.path.abspath(template_path)}")
        return os.path.abspath(template_path)

    except Exception as e:
        logger.error(f"Errore durante la creazione del template: {e}")
        logger.error(traceback.format_exc())
        return None


def crea_documento_denuncia(detentore_id: int, parent_widget=None) -> Tuple[bool, str]:
    """
    Crea il documento di denuncia per la detenzione di armi.

    Args:
        detentore_id: ID del detentore
        parent_widget: Widget genitore per i dialoghi

    Returns:
        Tuple (successo, messaggio)
    """
    if not detentore_id:
        logger.error("ID detentore non valido")
        QMessageBox.warning(parent_widget, "Attenzione", "ID detentore non valido.")
        return False, "ID detentore non valido"

    logger.info(f"Avvio generazione denuncia per ID_Detentore = {detentore_id}")
    progress = None

    try:
        # Mostra un dialogo di progresso
        progress = QProgressDialog("Generazione denuncia in corso...", "Annulla", 0, 5, parent_widget)
        progress.setWindowTitle("Generazione Denuncia")
        progress.setWindowModality(Qt.WindowModal)
        progress.setValue(0)
        progress.show()

        # Connessione al database
        with sqlite3.connect("gestione_armi.db") as conn:
            cursor = conn.cursor()

            # 1. Recupera i dati del detentore
            progress.setValue(1)
            detentore = get_detentore_data(cursor, detentore_id)
            if not detentore:
                progress.close()
                msg = f"Detentore con ID {detentore_id} non trovato nel database"
                QMessageBox.warning(parent_widget, "Attenzione", msg)
                return False, msg

            # 2. Recupera le armi
            progress.setValue(2)
            armi, armi_per_categoria = get_armi_data(cursor, detentore_id)

            if not armi:
                msg = f"Attenzione: nessuna arma trovata per il detentore {detentore['nome']} {detentore['cognome']}"
                logger.warning(msg)

                # Chiedi conferma per continuare senza armi
                progress.close()
                reply = QMessageBox.question(
                    parent_widget,
                    "Nessuna arma trovata",
                    f"Non sono state trovate armi per {detentore['nome']} {detentore['cognome']}.\n\nVuoi continuare comunque con la generazione della denuncia?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.No
                )

                if reply == QMessageBox.No:
                    return False, "Operazione annullata: nessuna arma trovata"

                # Ricrea il dialogo di progresso
                progress = QProgressDialog("Generazione denuncia in corso...", "Annulla", 0, 5, parent_widget)
                progress.setWindowTitle("Generazione Denuncia")
                progress.setWindowModality(Qt.WindowModal)
                progress.setValue(2)
                progress.show()

            # 3. Trova il template
            progress.setValue(3)
            template_candidates = [
                "denuncia_professionale_template.docx",
                "Denuncia_armi_completo.docx"
            ]

            template_path = None
            for template_name in template_candidates:
                template_path = find_template_path(template_name)
                if template_path:
                    break

            if not template_path:
                # Crea un template base se non ne esiste uno
                logger.info("Nessun template trovato, creazione di un template base")
                template_path = crea_template_professionale()

            if not template_path:
                progress.close()
                msg = f"Impossibile trovare o creare un template per la denuncia."
                QMessageBox.critical(parent_widget, "Errore", msg)
                return False, msg

            # 4. Prepara il contesto e genera il documento
            logger.info(f"Usando template: {template_path}")
            progress.setValue(4)

            context = prepare_context(detentore, armi, armi_per_categoria)

            # Debug del contesto per verificare la presenza delle armi
            logger.debug(f"Contesto per il template: {context}")
            logger.debug(f"Numero di armi nel contesto: {len(context['armi'])}")
            logger.debug(f"Numero di categorie nel contesto: {len(context['categorie'])}")

            # Carica e renderizza il template
            try:
                doc = DocxTemplate(template_path)
                doc.render(context)
            except jinja2.exceptions.UndefinedError as e:
                progress.close()
                logger.error(f"Errore durante la renderizzazione del template: {e}")
                logger.error(traceback.format_exc())

                # Chiedi all'utente se vuole usare un template base semplificato
                reply = QMessageBox.question(
                    parent_widget,
                    "Errore nel template",
                    f"Il template attuale ha un problema: {str(e)}.\n\nVuoi utilizzare un template base semplificato?",
                    QMessageBox.Yes | QMessageBox.No,
                    QMessageBox.Yes
                )

                if reply == QMessageBox.Yes:
                    try:
                        # Crea un template base
                        template_path = crea_template_professionale()
                        if template_path:
                            # Riprova con il nuovo template
                            doc = DocxTemplate(template_path)
                            doc.render(context)
                        else:
                            return False, "Impossibile creare un template base"
                    except Exception as new_e:
                        logger.error(f"Errore con il template base: {new_e}")
                        return False, f"Errore con il template base: {str(new_e)}"
                else:
                    return False, f"Errore durante la generazione del documento: {str(e)}"

            except Exception as e:
                progress.close()
                logger.error(f"Errore durante la renderizzazione del template: {e}")
                logger.error(traceback.format_exc())
                QMessageBox.critical(
                    parent_widget,
                    "Errore",
                    f"Errore durante la generazione del documento: {str(e)}"
                )
                return False, f"Errore durante la generazione del documento: {str(e)}"

            # 5. Salva il documento
            nome_file_default = f"Denuncia_Armi_{detentore['cognome']}_{detentore['nome']}_{datetime.now().strftime('%d_%m_%Y')}.docx"
            file_path, _ = QFileDialog.getSaveFileName(
                parent_widget,
                "Salva documento",
                nome_file_default,
                "Documenti Word (*.docx);;Tutti i file (*)"
            )

            if not file_path:
                progress.close()
                return False, "Operazione annullata dall'utente"

            progress.setValue(5)
            doc.save(file_path)
            logger.info(f"Documento salvato in: {file_path}")

            # Chiudi il dialogo di progresso
            progress.close()

            # Chiedi all'utente se vuole aprire il documento
            reply = QMessageBox.question(
                parent_widget,
                "Documento generato",
                f"Il documento è stato salvato con successo.\nVuoi aprire il documento?",
                QMessageBox.Yes | QMessageBox.No,
                QMessageBox.Yes
            )

            if reply == QMessageBox.Yes:
                logger.info("Apertura documento...")
                try:
                    if sys.platform == "win32":
                        os.startfile(file_path)
                    elif sys.platform == "darwin":  # macOS
                        subprocess.call(('open', file_path))
                    else:  # linux/unix
                        subprocess.call(('xdg-open', file_path))
                except Exception as e:
                    logger.error(f"Errore durante l'apertura del documento: {e}")
                    return True, f"Documento salvato in: {file_path}\nImpossibile aprire automaticamente il file: {str(e)}"

            return True, f"Documento salvato con successo in: {file_path}"

    except Exception as e:
        logger.error("Errore durante la generazione della denuncia")
        logger.error(traceback.format_exc())
        if progress and progress.isVisible():
            progress.close()
        QMessageBox.critical(parent_widget, "Errore Critico", f"Si è verificato un errore imprevisto:\n{str(e)}")
        return False, f"Si è verificato un errore imprevisto: {str(e)}"


def debug_db(detentore_id: int) -> None:
    """
    Funzione di debug per esaminare i dati del database relativi a un detentore.

    Args:
        detentore_id: ID del detentore
    """
    try:
        with sqlite3.connect("gestione_armi.db") as conn:
            cursor = conn.cursor()

            # Verifica i dati del detentore
            cursor.execute("SELECT * FROM detentori WHERE ID_Detentore = ?", (detentore_id,))
            detentore = cursor.fetchone()
            if detentore:
                print(f"Detentore trovato: ID={detentore_id}")

                # Verifica le armi
                cursor.execute("SELECT * FROM armi WHERE ID_Detentore = ?", (detentore_id,))
                armi = cursor.fetchall()
                print(f"Armi trovate: {len(armi)}")

                # Raggruppa le armi per categoria
                categorie = {}
                for arma in armi:
                    categoria = arma[7] if len(arma) > 7 and arma[7] else "ALTRA CATEGORIA"
                    if categoria not in categorie:
                        categorie[categoria] = []
                    categorie[categoria].append(arma)

                # Mostra le armi raggruppate per categoria
                print("\nArmi raggruppate per categoria:")
                for categoria, armi_cat in categorie.items():
                    print(f"\n{categoria} ({len(armi_cat)} armi):")
                    for i, arma in enumerate(armi_cat):
                        print(f"  {i + 1}. {arma[3]} {arma[4]} - Matricola: {arma[6]}")
            else:
                print(f"Nessun detentore trovato con ID {detentore_id}")
    except Exception as e:
        print(f"Errore durante il debug del database: {e}")


def test_connection():
    """Verifica la connessione al database e la struttura delle tabelle"""
    try:
        with sqlite3.connect("gestione_armi.db") as conn:
            cursor = conn.cursor()

            # Verifica la tabella detentori
            cursor.execute("PRAGMA table_info(detentori)")
            columns_detentori = [col[1] for col in cursor.fetchall()]
            required_detentori = ["ID_Detentore", "Nome", "Cognome", "CodiceFiscale"]

            for col in required_detentori:
                if col not in columns_detentori:
                    return False, f"Colonna {col} mancante nella tabella detentori"

            # Verifica la tabella armi
            cursor.execute("PRAGMA table_info(armi)")
            columns_armi = [col[1] for col in cursor.fetchall()]
            required_armi = ["ID_ArmaDetenuta", "ID_Detentore", "Matricola"]

            for col in required_armi:
                if col not in columns_armi:
                    return False, f"Colonna {col} mancante nella tabella armi"

            return True, "Connessione al database verificata con successo"
    except Exception as e:
        return False, f"Errore durante la verifica del database: {str(e)}"


if __name__ == "__main__":
    # Test standalone
    from PyQt5.QtWidgets import QApplication

    app = QApplication(sys.argv)

    # Verifica la connessione al database
    success, message = test_connection()
    print(message)

    if success and len(sys.argv) > 1:
        try:
            detentore_id = int(sys.argv[1])
            # Esegui debug del database
            print("Esecuzione debug database:")
            debug_db(detentore_id)

            # Genera la denuncia
            print("\nGenerazione denuncia:")
            success, message = crea_documento_denuncia(detentore_id)
            print(message)
        except ValueError:
            print("Argomento non valido. Usa: python GeneraDenuncia.py <id_detentore>")

    sys.exit(0)